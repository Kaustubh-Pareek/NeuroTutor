from docx import Document
import io
from flask import send_file

def export_to_word(data):
    summary = data.get("summary", "")
    questions = data.get("questions", "")
    concepts = data.get("concepts", "")
    filename = data.get("filename", "StudyGuide")

    doc = Document()
    doc.add_heading("Study Guide", level=0)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Key Concepts", level=1)
    doc.add_paragraph(concepts)

    doc.add_heading("Practice Questions", level=1)
    doc.add_paragraph(questions)

    buffer = io.BytesIO()# this creates an in-memory file -  a fake file stored in RAM
    doc.save(buffer)

    # moves the file pointer to the start of the buffer
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{filename}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
